#!/usr/bin/env python3
"""Build the frozen SharePoint Copilot pilot corpus and development cases."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import subprocess
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable
from urllib.parse import urlsplit
from xml.etree import ElementTree

from docx import Document
from docx.shared import Pt

from build_family_word import (
    EXPECTED_PROJECTION_SHA256,
    EXPECTED_RECORD_SCHEMA,
    SOURCE_COMMIT,
    SourceBinding,
    add_applicability,
    add_journeys,
    add_label_value,
    add_list_item,
    build_document,
    configure_document,
    create_numbering,
    read_governed_record,
    set_run_font,
)


EXPERIMENT_DIR = Path(__file__).resolve().parent
PROFILE_FULL = "word-full-v1"
PROFILE_RETRIEVAL = "word-retrieval-v2"
PROFILE_NAMES = (PROFILE_FULL, PROFILE_RETRIEVAL)
RETRIEVAL_TARGET_CHARACTERS = 20_000
RETRIEVAL_HARD_LIMIT_CHARACTERS = 36_000
FIXED_ZIP_TIMESTAMP = (1980, 1, 1, 0, 0, 0)

EXPECTED_PILOT_HTML_SHA256: dict[str, str] = {
    "apply-for-school-place": "4030a0d5dfb92ec4154bec3b4d57a07365deaa8052177135fe1abfd986a24205",
    "appeal-school-admission": "70ae0bf2082f1fd567b2e4aa75bd6f5174d9ab49bce3541cd1baa42d91a4fd32",
    "challenge-school-exclusion": "9d12ac2399ee7e85d0df48d19887b7d70ae2f9d22e5c75f8233d2b4edc876469",
    "apply-for-free-school-meals": "abdce3063083c397bca7fe18a0be116f18daa7e948e53302ec0ab7068a52fc13",
    "apply-for-social-housing": "bc5e56c6c14ff3f5987bb8b922a2f3bbf5a3c0d5c201877ef4a7ce3ec10c14c5",
    "get-homelessness-help": "9ec63d11f4297afe10e86c6584f4d4db046c1e4c189ea2a6a8c25b9ef81e2128",
    "rent-private-home": "1a1c1413addb85e16e76057333fbf1f87034dd61d86f6a92e28caa243833834b",
    "protect-tenancy-deposit": "7dd24ff45480ba69ef1d278440e5478a2eda5bcd15674fd14c804c55ccb889b7",
    "use-public-job-search": "cc16342644c1f0aa51a5a2c7a5b08788e02c9944bfd16f28945f84c9272f44a3",
    "get-jobcentre-support": "6ea2ce1d329177be2f2ddb77386c46fc8103b644c1f64d8e71f7e03f733d71ba",
    "get-redundancy-support": "cd0508d2b91d2b6b5fcfb6f89bcf6326b46bfc062a34fe1fc5baf59283ff337b",
    "claim-universal-credit-while-unemployed": "8af1646c4ef9ab1872a7225b05e0108c7d5b6e4308d4daf318b017f0177339e7",
    "register-a-death": "47c03ef4042870bc3e1ef141575e633ba0b162076f55d4779c92598e0b19c288",
    "arrange-funeral": "aca3bcdbf1c4d31801c4058b820488e5ce3473e6ef2227c04312dcb2c6f784c2",
    "arrange-burial-or-cremation": "24264f7266ff1b3e9a3b56dbb9332d5b7a7f4061abf3f476c96e887f7a11b195",
    "administer-an-estate": "2be13fcb04a08bb19f243fb3ee887f88359a0ec91b0a9b699389e8ab0441c003",
    "obtain-uk-passport": "eb793e726c84fcb3f460c818f600240793a9c14b581fab57997039233736f416",
    "renew-passport": "72cfbe85483aacd3ac0c604f80b2caaafcb18300b346b89d0071c880eeef8a0a",
    "apply-for-visa-or-immigration-permission": "4b78713a7f2430dbdd120d30d65677abdc8726cbc764f57b4fc4bff47c57becb",
    "apply-for-citizenship": "063c9cc7672a5de02db3c0dfaeaca6f0ff21f978ab2fe6daa8d7a4eefdb22ed2",
}

PILOT_GROUPS: dict[str, tuple[str, ...]] = {
    "school": (
        "apply-for-school-place",
        "appeal-school-admission",
        "challenge-school-exclusion",
        "apply-for-free-school-meals",
    ),
    "housing": (
        "apply-for-social-housing",
        "get-homelessness-help",
        "rent-private-home",
        "protect-tenancy-deposit",
    ),
    "work": (
        "use-public-job-search",
        "get-jobcentre-support",
        "get-redundancy-support",
        "claim-universal-credit-while-unemployed",
    ),
    "bereavement": (
        "register-a-death",
        "arrange-funeral",
        "arrange-burial-or-cremation",
        "administer-an-estate",
    ),
    "passport-and-immigration": (
        "obtain-uk-passport",
        "renew-passport",
        "apply-for-visa-or-immigration-permission",
        "apply-for-citizenship",
    ),
}

AMBIGUOUS_PAIRS: tuple[tuple[str, str], ...] = (
    ("apply-for-school-place", "appeal-school-admission"),
    ("apply-for-social-housing", "get-homelessness-help"),
    ("get-jobcentre-support", "claim-universal-credit-while-unemployed"),
    ("obtain-uk-passport", "renew-passport"),
)

NEGATIVE_FAMILIES: tuple[str, ...] = (
    "access-dental-care",
    "claim-child-benefit",
    "get-blue-badge",
    "register-a-birth",
)


def pilot_family_ids() -> tuple[str, ...]:
    return tuple(family_id for group in PILOT_GROUPS.values() for family_id in group)


def parse_args() -> argparse.Namespace:
    default_source_repo = EXPERIMENT_DIR.parents[2] / "okf-uk-living"
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-repo", type=Path, default=default_source_repo)
    parser.add_argument(
        "--check",
        action="store_true",
        help="Verify existing artefacts against the manifest without rebuilding.",
    )
    return parser.parse_args()


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def write_text_if_changed(path: Path, text: str) -> None:
    encoded = text.encode("utf-8")
    if path.exists() and path.read_bytes() == encoded:
        return
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(encoded)


def normalise_docx_package(path: Path) -> None:
    """Make ZIP container ordering and timestamps deterministic."""

    with zipfile.ZipFile(path, "r") as source:
        entries = [(name, source.read(name)) for name in sorted(source.namelist())]

    with tempfile.NamedTemporaryFile(
        dir=path.parent,
        prefix=f".{path.stem}-",
        suffix=".docx",
        delete=False,
    ) as temporary:
        temporary_path = Path(temporary.name)

    try:
        with zipfile.ZipFile(
            temporary_path,
            "w",
            compression=zipfile.ZIP_DEFLATED,
            compresslevel=9,
        ) as target:
            for name, data in entries:
                info = zipfile.ZipInfo(name, FIXED_ZIP_TIMESTAMP)
                info.compress_type = zipfile.ZIP_DEFLATED
                info.external_attr = 0o600 << 16
                target.writestr(info, data)
        temporary_path.replace(path)
    finally:
        if temporary_path.exists():
            temporary_path.unlink()


def document_text(path: Path) -> str:
    with zipfile.ZipFile(path, "r") as package:
        root = ElementTree.fromstring(package.read("word/document.xml"))
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    return "".join(node.text or "" for node in root.iter(f"{namespace}t"))


def validate_web_url(url: str, *, label: str) -> None:
    if any(character.isspace() for character in url):
        raise ValueError(f"Whitespace is not allowed in {label}: {url!r}")
    parsed = urlsplit(url)
    if parsed.scheme not in {"http", "https"} or not parsed.hostname:
        raise ValueError(f"Expected a credential-free HTTP(S) URL for {label}: {url!r}")
    if parsed.username is not None or parsed.password is not None:
        raise ValueError(f"Credentials are not allowed in {label}: {url!r}")
    try:
        port = parsed.port
    except ValueError as error:
        raise ValueError(f"Invalid port in {label}: {url!r}") from error
    if port is not None and not 1 <= port <= 65_535:
        raise ValueError(f"Port outside 1–65535 in {label}: {url!r}")


def validate_ordered_items(
    items: list[dict[str, Any]],
    *,
    label: str,
) -> None:
    orders = [item["order"] for item in items]
    if sorted(orders) != list(range(len(items))):
        raise ValueError(f"{label} orders are not contiguous from zero: {orders}")


def validate_governed_record(record: dict[str, Any], binding: SourceBinding) -> None:
    family = record["family"]
    sources = family["sources"]
    source_ids = [source["id"] for source in sources]
    if len(source_ids) != len(set(source_ids)):
        raise ValueError(f"Duplicate source IDs in {binding.family_id}")
    positions = [source["position"] for source in sources]
    if sorted(positions) != list(range(len(sources))):
        raise ValueError(
            f"Source positions are not contiguous from zero in {binding.family_id}: {positions}"
        )
    source_id_set = set(source_ids)
    for source in sources:
        validate_web_url(source["url"], label=f"{binding.family_id} source {source['id']}")
    validate_web_url(binding.public_url, label=f"{binding.family_id} public record")

    episodes = family["episodes"]
    validate_ordered_items(episodes, label=f"{binding.family_id} episodes")
    episode_ids = [episode["id"] for episode in episodes]
    if len(episode_ids) != len(set(episode_ids)):
        raise ValueError(f"Duplicate episode IDs in {binding.family_id}")
    for episode in episodes:
        validate_ordered_items(
            episode["steps"],
            label=f"{binding.family_id} episode {episode['id']} steps",
        )
        step_ids = [step["id"] for step in episode["steps"]]
        if len(step_ids) != len(set(step_ids)):
            raise ValueError(
                f"Duplicate step IDs in {binding.family_id} episode {episode['id']}"
            )
        for step in episode["steps"]:
            unresolved = sorted(set(step["source_ids"]) - source_id_set)
            if unresolved:
                raise ValueError(
                    f"Unresolved source IDs in {binding.family_id} step {step['id']}: "
                    f"{', '.join(unresolved)}"
                )

    for applicability in family["applicability"]:
        for variant in applicability["route_variants"]:
            if variant["primary_source"] not in source_id_set:
                raise ValueError(
                    f"Unresolved primary source in {binding.family_id} route "
                    f"{variant['id']}: {variant['primary_source']}"
                )


def verify_projection_blob(
    source_repo: Path,
    source_ref: str,
    record: dict[str, Any],
) -> None:
    projection = record["source_projection"]
    completed = subprocess.run(
        [
            "git",
            "-C",
            str(source_repo),
            "show",
            f"{source_ref}:{projection['path']}",
        ],
        check=True,
        capture_output=True,
    )
    content = completed.stdout
    if len(content) != projection["bytes"]:
        raise ValueError(
            "Journey projection byte count does not match the governed record: "
            f"expected {projection['bytes']}, got {len(content)}"
        )
    actual_sha = hashlib.sha256(content).hexdigest()
    if actual_sha != projection["sha256"]:
        raise ValueError(
            "Journey projection SHA-256 does not match the governed record: "
            f"expected {projection['sha256']}, got {actual_sha}"
        )


def assert_tokens_in_order(text: str, tokens: list[str], *, label: str) -> None:
    cursor = 0
    for token in tokens:
        position = text.find(token, cursor)
        if position < 0:
            raise ValueError(f"Missing or out-of-order token in {label}: {token!r}")
        cursor = position + len(token)


def retrieval_tokens(record: dict[str, Any], binding: SourceBinding) -> list[str]:
    family = record["family"]
    tokens = [
        family["title"],
        family["id"],
        family["description"],
        family["interaction_boundary"],
        binding.html_sha256,
        binding.record_sha256,
        record["schema"],
        record["source_projection"]["sha256"],
        *family["aliases"],
        *family["situations"],
        *family["user_needs"],
        family["route"],
        family["domain"]["id"],
        family["process"]["id"],
        family["status"],
        family["assertion_status"],
    ]
    for applicability in family["applicability"]:
        tokens.append(applicability["jurisdiction"])
        tokens.append(applicability["state"])
        for variant in applicability["route_variants"]:
            tokens.extend(
                [
                    variant["id"],
                    variant["provider"],
                    variant["primary_source"],
                ]
            )
    for episode in sorted(family["episodes"], key=lambda item: item["order"]):
        tokens.extend([episode["title"], episode["kind"], episode["id"]])
        for step in sorted(episode["steps"], key=lambda item: item["order"]):
            tokens.extend([step["interaction"], step["id"], *step["source_ids"]])
    for source in sorted(family["sources"], key=lambda item: item["position"]):
        tokens.extend([source["title"], source["url"], source["id"]])
    tokens.extend(
        [
            family["review"]["population_gate"],
            family["review"]["specialist_review"],
            *family["limitations"],
        ]
    )
    for related in family["related_families"]:
        tokens.extend([related["title"], related["id"], related["relationship"]])
    return tokens


def verify_full_appendix(path: Path, record_text: str) -> None:
    doc = Document(path)
    paragraphs = doc.paragraphs
    heading_index = next(
        (
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph.text == "Complete governed record (technical appendix)"
        ),
        None,
    )
    if heading_index is None:
        raise ValueError(f"Technical appendix heading is missing from {path}")
    readable_text = "\n".join(
        paragraph.text for paragraph in paragraphs[:heading_index]
    )
    if re.search(r"\[[^]]+]\([^)]+\)", readable_text):
        raise ValueError(f"Unrendered Markdown link in readable content of {path}")
    extracted = "\n".join(paragraph.text for paragraph in paragraphs[heading_index + 2 :])
    if extracted != record_text:
        raise ValueError(f"Technical appendix does not exactly match the source JSON in {path}")


def verify_retrieval_structure(
    path: Path,
    record: dict[str, Any],
    binding: SourceBinding,
) -> None:
    doc = Document(path)
    if doc.tables:
        raise ValueError(f"Retrieval profile must not contain tables: {path}")
    with zipfile.ZipFile(path, "r") as package:
        document_xml = package.read("word/document.xml")
        if b"<w:txbxContent" in document_xml:
            raise ValueError(f"Retrieval profile must not contain text boxes: {path}")
        document_root = ElementTree.fromstring(document_xml)
        relationships_root = ElementTree.fromstring(
            package.read("word/_rels/document.xml.rels")
        )
    relationship_namespace = "{http://schemas.openxmlformats.org/package/2006/relationships}"
    relationship_targets = {
        relationship.attrib["Id"]: relationship.attrib["Target"]
        for relationship in relationships_root.findall(
            f"{relationship_namespace}Relationship"
        )
        if relationship.attrib.get("TargetMode") == "External"
    }
    word_namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    relationship_id_attribute = (
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
    )
    visible_hyperlinks: list[tuple[str, str]] = []
    for hyperlink in document_root.iter(f"{word_namespace}hyperlink"):
        relationship_id = hyperlink.attrib.get(relationship_id_attribute)
        if not relationship_id or relationship_id not in relationship_targets:
            continue
        text = "".join(
            node.text or "" for node in hyperlink.iter(f"{word_namespace}t")
        )
        visible_hyperlinks.append((relationship_targets[relationship_id], text))
    visible_text = document_text(path)
    assert_tokens_in_order(
        visible_text,
        retrieval_tokens(record, binding),
        label=str(path),
    )
    for source in record["family"]["sources"]:
        if source["url"] not in visible_text:
            raise ValueError(f"Official URL is not visible in {path}: {source['url']}")
        if (source["url"], source["url"]) not in visible_hyperlinks:
            raise ValueError(
                f"Official hyperlink target and visible text differ in {path}: "
                f"{source['url']}"
            )


def add_plain_note(doc: Document, label: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    label_run = paragraph.add_run(f"{label} ")
    set_run_font(label_run, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run)


def add_retrieval_title(
    doc: Document,
    family: dict[str, Any],
    binding: SourceBinding,
) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(3)
    kicker_run = kicker.add_run("EXPLORE OKF — RETRIEVAL PROFILE V2")
    set_run_font(kicker_run, size=9.5, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    title.paragraph_format.keep_with_next = True
    title_run = title.add_run(family["title"])
    set_run_font(title_run, size=24, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle_run = subtitle.add_run(
        "Concise governed family record for SharePoint and AI discovery"
    )
    set_run_font(subtitle_run, size=12)

    add_label_value(doc, "Stable family ID", family["id"], code=True)
    add_label_value(doc, "Description", family["description"])
    add_label_value(doc, "Interaction boundary", family["interaction_boundary"])
    add_label_value(doc, "Unique family HTML SHA-256", binding.html_sha256, code=True)
    add_label_value(doc, "Governed record SHA-256", binding.record_sha256, code=True)

    add_plain_note(
        doc,
        "Exploratory.",
        "This deterministic derivative is a discovery aid, not an official service "
        "or current or personalised advice. Check the current official source before acting.",
    )


def add_retrieval_signals(doc: Document, family: dict[str, Any]) -> None:
    doc.add_heading("When this family may match", level=1)

    doc.add_heading("Authored terms", level=2)
    bullet_id = create_numbering(doc, bullet=True)
    for alias in family["aliases"]:
        add_list_item(doc, alias, bullet_id)

    doc.add_heading("Authored example situations", level=2)
    bullet_id = create_numbering(doc, bullet=True)
    for situation in family["situations"]:
        add_list_item(doc, situation, bullet_id)

    doc.add_heading("Authored user needs", level=2)
    bullet_id = create_numbering(doc, bullet=True)
    for need in family["user_needs"]:
        add_list_item(doc, need, bullet_id)


def add_retrieval_provenance(
    doc: Document,
    record: dict[str, Any],
    binding: SourceBinding,
) -> None:
    doc.add_heading("Retrieval and provenance gate", level=1)
    add_label_value(doc, "Record schema", record["schema"], code=True)
    add_label_value(
        doc,
        "Source projection SHA-256",
        record["source_projection"]["sha256"],
        code=True,
    )
    add_label_value(doc, "Source projection path", record["source_projection"]["path"], code=True)
    add_label_value(doc, "Source projection bytes", record["source_projection"]["bytes"], code=True)
    add_label_value(doc, "Snapshot", record["snapshot"], code=True)
    add_label_value(doc, "Generated at", record["generated_at"], code=True)
    add_label_value(doc, "Source Git commit", binding.source_ref, code=True)
    add_label_value(doc, "Source family path", binding.source_path, code=True)
    add_label_value(
        doc,
        "Public governed record",
        binding.public_url,
        url=binding.public_url,
    )
    add_plain_note(
        doc,
        "Gatekeeping rule.",
        "Report the record schema and source projection SHA-256 before using the "
        "family content. If either is unavailable, stop rather than infer details.",
    )


def add_family_placement(doc: Document, family: dict[str, Any]) -> None:
    doc.add_heading("Family placement and status", level=1)
    add_label_value(doc, "Family route", family["route"], code=True)
    add_label_value(doc, "Domain", f"{family['domain']['title']} ({family['domain']['id']})")
    add_label_value(doc, "Domain route", family["domain"]["route"], code=True)
    add_label_value(
        doc,
        "Enclosing process",
        f"{family['process']['title']} ({family['process']['id']})",
    )
    add_label_value(doc, "Process route", family["process"]["route"], code=True)
    add_label_value(doc, "Population status", family["status"], code=True)
    add_label_value(doc, "Assertion status", family["assertion_status"], code=True)


def add_linear_sources(doc: Document, family: dict[str, Any]) -> None:
    doc.add_heading("Official source handoffs", level=1)
    doc.add_paragraph(
        "These links locate current official services. Check the relevant official source before acting."
    )
    for source in sorted(family["sources"], key=lambda item: item["position"]):
        jurisdiction = ", ".join(source["jurisdictions"]) or "Jurisdiction not stated"
        doc.add_heading(f"{jurisdiction}: {source['title']}", level=2)
        add_label_value(doc, "Official URL", source["url"], url=source["url"])
        add_label_value(doc, "Owner", source["owner"])
        add_label_value(doc, "Source ID", source["id"], code=True)
        add_label_value(doc, "Source route", source["route"], code=True)
        add_label_value(doc, "Authority role", source["authority_role"], code=True)
        add_label_value(doc, "Summary", source["summary"])
        add_label_value(doc, "Observed at", source["observed_at"], code=True)
        add_label_value(doc, "Freshness", source["freshness"], code=True)
        add_label_value(doc, "Jurisdiction basis", source["jurisdiction_basis"], code=True)
        add_label_value(doc, "Rights decision", source["rights_decision"], code=True)
        if source["limitations"]:
            doc.add_heading("Source limitations", level=3)
            bullet_id = create_numbering(doc, bullet=True)
            for limitation in source["limitations"]:
                add_list_item(doc, limitation, bullet_id)


def add_linear_review_and_related(doc: Document, family: dict[str, Any]) -> None:
    doc.add_heading("Review status and limitations", level=1)
    add_label_value(doc, "Population gate", family["review"]["population_gate"], code=True)
    add_label_value(doc, "Specialist review", family["review"]["specialist_review"], code=True)
    add_plain_note(
        doc,
        "Important distinction.",
        "Population-complete means the family is present for discovery. It does not "
        "mean that the record has passed specialist review or that operational claims "
        "are release-grade.",
    )
    bullet_id = create_numbering(doc, bullet=True)
    for limitation in family["limitations"]:
        add_list_item(doc, limitation, bullet_id)

    doc.add_heading("Related families", level=1)
    if not family["related_families"]:
        doc.add_paragraph("No related families are stated in this record.")
        return
    for related in family["related_families"]:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        label_run = paragraph.add_run("Related family: ")
        set_run_font(label_run, bold=True)
        detail_run = paragraph.add_run(
            f"{related['title']} — {related['id']}; relationship: "
            f"{related['relationship']}; sequenced: "
            f"{str(related['sequenced']).lower()}"
        )
        set_run_font(detail_run)


def build_retrieval_document(
    record: dict[str, Any],
    binding: SourceBinding,
    output: Path,
) -> None:
    doc = Document()
    configure_document(doc, header_label="Retrieval profile v2")
    family = record["family"]

    doc.core_properties.title = f"{family['title']} — governed family record"
    doc.core_properties.subject = "SharePoint and Microsoft 365 Copilot retrieval experiment"
    doc.core_properties.author = "OKF Explorer"
    doc.core_properties.last_modified_by = "OKF Explorer"
    doc.core_properties.keywords = (
        "OKF, Open Knowledge Format, SharePoint, Microsoft 365 Copilot, "
        f"{family['id']}, {', '.join(family['aliases'])}"
    )
    doc.core_properties.comments = (
        "Deterministically generated from the pinned OKF UK Living governed family record."
    )
    source_time = datetime.fromisoformat(record["generated_at"].replace("Z", "+00:00"))
    doc.core_properties.created = source_time
    doc.core_properties.modified = source_time
    doc.core_properties.revision = 1

    add_retrieval_title(doc, family, binding)
    add_retrieval_provenance(doc, record, binding)
    add_retrieval_signals(doc, family)
    add_family_placement(doc, family)
    add_applicability(doc, family)
    add_journeys(doc, family)
    add_linear_sources(doc, family)
    add_linear_review_and_related(doc, family)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def select_indirect_signal(family: dict[str, Any]) -> tuple[str, str]:
    normalised_title = family["title"].casefold()
    for alias in family["aliases"]:
        if alias.casefold() != normalised_title:
            return "alias", alias
    if len(family["situations"]) > 1:
        return "situation", family["situations"][1]
    return "description", family["description"]


def expected_identity(
    record: dict[str, Any],
    binding: SourceBinding,
) -> dict[str, str]:
    return {
        "family_id": record["family"]["id"],
        "family_title": record["family"]["title"],
        "family_html_sha256": binding.html_sha256,
        "governed_record_sha256": binding.record_sha256,
        "record_schema": record["schema"],
        "source_projection_sha256": record["source_projection"]["sha256"],
    }


def make_development_cases(
    records: dict[str, tuple[str, dict[str, Any], SourceBinding]],
    negative_records: dict[str, tuple[str, dict[str, Any], SourceBinding]],
) -> list[dict[str, Any]]:
    cases: list[dict[str, Any]] = []
    for position, family_id in enumerate(pilot_family_ids(), start=1):
        _, record, binding = records[family_id]
        family = record["family"]
        authored_situation = family["situations"][0]
        cases.append(
            {
                "schema": "explore-okf-development-case.v1",
                "case_id": f"dev-positive-{position:02d}-situation",
                "case_kind": "clear_authored_situation",
                "situation": authored_situation,
                "expected_behaviour": "select_one_family",
                "expected": expected_identity(record, binding),
                "evidence": {
                    "family_id": family_id,
                    "field": "family.situations[0]",
                    "text": authored_situation,
                },
                "synthetic_personal_data": False,
                "development_only": True,
            }
        )

        signal_kind, signal = select_indirect_signal(family)
        cases.append(
            {
                "schema": "explore-okf-development-case.v1",
                "case_id": f"dev-positive-{position:02d}-indirect",
                "case_kind": "governed_indirect_signal",
                "situation": f"I need help with this situation: {signal}.",
                "expected_behaviour": "select_one_family",
                "expected": expected_identity(record, binding),
                "evidence": {
                    "family_id": family_id,
                    "field": f"family.{signal_kind}",
                    "text": signal,
                },
                "synthetic_personal_data": False,
                "development_only": True,
            }
        )

    for position, (left_id, right_id) in enumerate(AMBIGUOUS_PAIRS, start=1):
        _, left, _ = records[left_id]
        _, right, _ = records[right_id]
        left_situation = left["family"]["situations"][0]
        right_situation = right["family"]["situations"][0]
        cases.append(
            {
                "schema": "explore-okf-development-case.v1",
                "case_id": f"dev-ambiguous-{position:02d}",
                "case_kind": "deliberately_ambiguous_near_neighbours",
                "situation": (
                    f"My situation could involve either '{left_situation}' or "
                    f"'{right_situation}', but I have not said which outcome I need."
                ),
                "expected_behaviour": "clarify_without_selecting",
                "acceptable_family_ids": [left_id, right_id],
                "evidence": [
                    {
                        "family_id": left_id,
                        "field": "family.situations[0]",
                        "text": left_situation,
                    },
                    {
                        "family_id": right_id,
                        "field": "family.situations[0]",
                        "text": right_situation,
                    },
                ],
                "synthetic_personal_data": False,
                "development_only": True,
            }
        )

    for position, family_id in enumerate(NEGATIVE_FAMILIES, start=1):
        _, record, binding = negative_records[family_id]
        situation = record["family"]["situations"][0]
        cases.append(
            {
                "schema": "explore-okf-development-case.v1",
                "case_id": f"dev-closed-corpus-negative-{position:02d}",
                "case_kind": "closed_corpus_negative",
                "situation": situation,
                "expected_behaviour": "say_not_covered",
                "excluded_true_family": expected_identity(record, binding),
                "evidence": {
                    "family_id": family_id,
                    "field": "family.situations[0]",
                    "text": situation,
                },
                "synthetic_personal_data": False,
                "development_only": True,
            }
        )

    if len(cases) != 48:
        raise AssertionError(f"Expected 48 development cases, generated {len(cases)}")
    return cases


def serialise_jsonl(items: Iterable[dict[str, Any]]) -> str:
    return "".join(
        f"{json.dumps(item, ensure_ascii=False, sort_keys=True)}\n" for item in items
    )


def build_profiles(
    source_repo: Path,
    source_ref: str,
    output_root: Path,
    selected_profiles: tuple[str, ...],
    selected_families: tuple[str, ...],
) -> tuple[
    dict[str, tuple[str, dict[str, Any], SourceBinding]],
    dict[str, list[dict[str, Any]]],
]:
    records: dict[str, tuple[str, dict[str, Any], SourceBinding]] = {}
    profile_entries: dict[str, list[dict[str, Any]]] = {
        profile: [] for profile in selected_profiles
    }
    projection_verified = False

    for family_id in selected_families:
        record_text, record, binding = read_governed_record(
            source_repo,
            source_ref,
            family_id,
        )
        expected_html_sha = EXPECTED_PILOT_HTML_SHA256[family_id]
        if binding.html_sha256 != expected_html_sha:
            raise ValueError(
                f"Pinned HTML digest mismatch for {family_id}: "
                f"expected {expected_html_sha}, got {binding.html_sha256}"
            )
        validate_governed_record(record, binding)
        if not projection_verified:
            verify_projection_blob(source_repo, source_ref, record)
            projection_verified = True
        records[family_id] = (record_text, record, binding)
        for profile in selected_profiles:
            output = output_root / profile / f"{family_id}.docx"
            active_locks = sorted(output.parent.glob("~$*.docx")) if output.parent.exists() else []
            if active_locks:
                raise RuntimeError(
                    f"Close Word before rebuilding {profile}; active lock files: "
                    f"{', '.join(str(path) for path in active_locks)}"
                )
            if profile == PROFILE_FULL:
                build_document(record_text, record, binding, output)
            elif profile == PROFILE_RETRIEVAL:
                build_retrieval_document(record, binding, output)
            else:
                raise AssertionError(f"Unexpected profile: {profile}")
            normalise_docx_package(output)
            visible_text = document_text(output)
            text_characters = len(visible_text)
            if profile == PROFILE_RETRIEVAL and text_characters > RETRIEVAL_TARGET_CHARACTERS:
                raise ValueError(
                    f"{family_id} retrieval profile contains {text_characters} "
                    f"characters; experiment target is {RETRIEVAL_TARGET_CHARACTERS}"
                )
            with zipfile.ZipFile(output, "r") as package:
                corrupt_member = package.testzip()
            if corrupt_member is not None:
                raise ValueError(f"Corrupt DOCX member in {output}: {corrupt_member}")
            if profile == PROFILE_FULL:
                verify_full_appendix(output, record_text)
            else:
                verify_retrieval_structure(output, record, binding)
            profile_entries[profile].append(
                {
                    "family_id": family_id,
                    "path": str(output.relative_to(EXPERIMENT_DIR)),
                    "sha256": sha256_file(output),
                    "bytes": output.stat().st_size,
                    "text_characters": text_characters,
                    "body_text_sha256": hashlib.sha256(
                        visible_text.encode("utf-8")
                    ).hexdigest(),
                    "within_retrieval_target": (
                        text_characters <= RETRIEVAL_TARGET_CHARACTERS
                        if profile == PROFILE_RETRIEVAL
                        else None
                    ),
                }
            )
    return records, profile_entries


def build_manifest(
    records: dict[str, tuple[str, dict[str, Any], SourceBinding]],
    profile_entries: dict[str, list[dict[str, Any]]],
    cases_output: Path,
    case_count: int,
) -> dict[str, Any]:
    families: list[dict[str, Any]] = []
    for group, family_ids in PILOT_GROUPS.items():
        for family_id in family_ids:
            if family_id not in records:
                continue
            _, record, binding = records[family_id]
            family = record["family"]
            families.append(
                {
                    "group": group,
                    "id": family_id,
                    "title": family["title"],
                    "domain_id": family["domain"]["id"],
                    "aliases": family["aliases"],
                    "jurisdictions": [
                        item["jurisdiction"] for item in family["applicability"]
                    ],
                    "population_status": family["status"],
                    "specialist_review": family["review"]["specialist_review"],
                    "source_path": binding.source_path,
                    "public_url": binding.public_url,
                    "family_html_sha256": binding.html_sha256,
                    "governed_record_sha256": binding.record_sha256,
                }
            )

    generated_values = sorted(
        {record[1]["generated_at"] for record in records.values()}
    )
    first_record = next(iter(records.values()))[1]
    projection = first_record["source_projection"]
    return {
        "schema": "explore-okf-sharepoint-pilot-corpus.v1",
        "purpose": (
            "Frozen challenge corpus for SharePoint and OneNote natural-language "
            "family discovery experiments"
        ),
        "source": {
            "repository": "okf-uk-living",
            "commit": SOURCE_COMMIT,
            "record_schema": EXPECTED_RECORD_SCHEMA,
            "source_projection_path": projection["path"],
            "source_projection_bytes": projection["bytes"],
            "source_projection_sha256": EXPECTED_PROJECTION_SHA256,
            "record_generated_at_values": generated_values,
        },
        "design": {
            "preset": "compact_reference_guide",
            "header_pattern": "customer_pack",
            "header_override": "linear_retrieval_title",
            "retrieval_target_characters": RETRIEVAL_TARGET_CHARACTERS,
            "retrieval_hard_limit_characters": RETRIEVAL_HARD_LIMIT_CHARACTERS,
        },
        "digest_definitions": {
            "family_html_sha256": "SHA-256 of the exact family HTML bytes at the pinned commit",
            "governed_record_sha256": (
                "SHA-256 of the exact UTF-8 unescaped JSON text embedded in the family HTML"
            ),
            "source_projection_sha256": (
                "SHA-256 stated by the governed record for the shared journey projection"
            ),
        },
        "corpus_kind": "deliberately_confusable_challenge_sample",
        "corpus_size": len(families),
        "build_status": (
            "complete" if len(families) == len(pilot_family_ids()) else "diagnostic_partial"
        ),
        "families": families,
        "profiles": {
            profile: {
                "document_count": len(entries),
                "documents": entries,
            }
            for profile, entries in profile_entries.items()
        },
        "development_cases": {
            "path": str(cases_output.relative_to(EXPERIMENT_DIR)),
            "count": case_count,
            "sha256": sha256_file(cases_output),
            "status": "development_only_not_final_holdout",
        },
        "publication": {
            "status": "private_experiment_not_release_grade",
            "bundle_rebuild_required": False,
            "official_advice": False,
        },
    }


def verify_manifest(manifest_path: Path) -> None:
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    if manifest["schema"] != "explore-okf-sharepoint-pilot-corpus.v1":
        raise ValueError(f"Unexpected manifest schema: {manifest['schema']!r}")
    if manifest["source"]["commit"] != SOURCE_COMMIT:
        raise ValueError("Manifest source commit does not match the pinned experiment")
    if manifest["source"]["source_projection_sha256"] != EXPECTED_PROJECTION_SHA256:
        raise ValueError("Manifest source projection digest does not match")

    checked = 0
    for profile_name, profile in manifest["profiles"].items():
        listed_paths = {
            str((EXPERIMENT_DIR / item["path"]).resolve())
            for item in profile["documents"]
        }
        profile_dir = EXPERIMENT_DIR / "profiles" / profile_name
        actual_paths = {
            str(path.resolve()) for path in profile_dir.glob("*.docx")
        }
        if listed_paths != actual_paths:
            missing = sorted(listed_paths - actual_paths)
            unmanifested = sorted(actual_paths - listed_paths)
            raise ValueError(
                f"Profile {profile_name} does not match its manifest; "
                f"missing={missing}, unmanifested={unmanifested}"
            )
        for item in profile["documents"]:
            path = EXPERIMENT_DIR / item["path"]
            if not path.is_file():
                raise FileNotFoundError(path)
            actual_sha = sha256_file(path)
            if actual_sha != item["sha256"]:
                raise ValueError(f"DOCX digest mismatch for {path}")
            visible_text = document_text(path)
            if len(visible_text) != item["text_characters"]:
                raise ValueError(f"DOCX character-count mismatch for {path}")
            if hashlib.sha256(visible_text.encode("utf-8")).hexdigest() != item[
                "body_text_sha256"
            ]:
                raise ValueError(f"DOCX body-text digest mismatch for {path}")
            with zipfile.ZipFile(path, "r") as package:
                corrupt_member = package.testzip()
            if corrupt_member is not None:
                raise ValueError(f"Corrupt DOCX member in {path}: {corrupt_member}")
            checked += 1

    cases = EXPERIMENT_DIR / manifest["development_cases"]["path"]
    if sha256_file(cases) != manifest["development_cases"]["sha256"]:
        raise ValueError("Development-case digest mismatch")
    case_lines = [
        line for line in cases.read_text(encoding="utf-8").splitlines() if line
    ]
    if len(case_lines) != manifest["development_cases"]["count"]:
        raise ValueError("Development-case count does not match the manifest")
    for line in case_lines:
        case = json.loads(line)
        if case["schema"] != "explore-okf-development-case.v1":
            raise ValueError(f"Unexpected development-case schema: {case['schema']!r}")
    print(f"Verified {checked} Word documents and the development cases")


def main() -> None:
    args = parse_args()
    output_root = EXPERIMENT_DIR / "profiles"
    manifest_output = EXPERIMENT_DIR / "corpus-manifest.json"
    cases_output = EXPERIMENT_DIR / "development-cases.v1.jsonl"
    if args.check:
        verify_manifest(manifest_output)
        return

    records, profile_entries = build_profiles(
        args.source_repo,
        SOURCE_COMMIT,
        output_root,
        PROFILE_NAMES,
        pilot_family_ids(),
    )

    negative_records = {
        family_id: read_governed_record(args.source_repo, SOURCE_COMMIT, family_id)
        for family_id in NEGATIVE_FAMILIES
    }
    cases = make_development_cases(records, negative_records)
    write_text_if_changed(cases_output, serialise_jsonl(cases))

    manifest = build_manifest(
        records,
        profile_entries,
        cases_output,
        len(cases),
    )
    write_text_if_changed(
        manifest_output,
        f"{json.dumps(manifest, ensure_ascii=False, indent=2, sort_keys=True)}\n",
    )
    verify_manifest(manifest_output)
    print(manifest_output)


if __name__ == "__main__":
    main()

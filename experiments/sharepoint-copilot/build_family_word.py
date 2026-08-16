#!/usr/bin/env python3
"""Build the first SharePoint/M365 Copilot OKF family Word experiment."""

from __future__ import annotations

import argparse
import hashlib
import html
import json
import re
import subprocess
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any
from urllib.parse import urljoin, urlsplit

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


SOURCE_COMMIT = "736d7dc4dbb4e44082f6b7786dd88afd55954792"
SOURCE_PATH_TEMPLATE = "explore/ai/families/{family_id}.html"
SOURCE_URL_ROOT = "https://chris-page-gov.github.io/okf-uk-living/"
EXPECTED_SCHOOL_HTML_SHA256 = (
    "4030a0d5dfb92ec4154bec3b4d57a07365deaa8052177135fe1abfd986a24205"
)
EXPECTED_RECORD_SCHEMA = "explore-okf-ai-family-record.v1"
EXPECTED_PROJECTION_SHA256 = (
    "646157327f3181bbef544613e8cd7398328c155dfb6939fcb9a3f1c883e07184"
)

INK = RGBColor(23, 33, 43)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(70, 87, 102)
TEAL = RGBColor(0, 107, 98)
WHITE = RGBColor(255, 255, 255)

DESIGN_PRESET = "compact_reference_guide"
HEADER_PATTERN = "customer_pack"
TITLE_BLOCK_OVERRIDE = "linear_retrieval_title"
MARKDOWN_LINK_PATTERN = re.compile(r"\[([^]]+)]\(([^)]+)\)")


@dataclass(frozen=True)
class SourceBinding:
    family_id: str
    source_path: str
    public_url: str
    source_ref: str
    html_sha256: str
    record_sha256: str


def parse_args() -> argparse.Namespace:
    repo_root = Path(__file__).resolve().parents[2]
    default_source_repo = repo_root.parent / "okf-uk-living"
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-repo", type=Path, default=default_source_repo)
    parser.add_argument("--source-ref", default=SOURCE_COMMIT)
    parser.add_argument("--family", default="apply-for-school-place")
    parser.add_argument("--output", type=Path)
    return parser.parse_args()


def read_governed_record(
    source_repo: Path,
    source_ref: str,
    family_id: str = "apply-for-school-place",
) -> tuple[str, dict[str, Any], SourceBinding]:
    if not re.fullmatch(r"[a-z0-9]+(?:-[a-z0-9]+)*", family_id):
        raise ValueError(f"Invalid family ID: {family_id!r}")
    source_path = SOURCE_PATH_TEMPLATE.format(family_id=family_id)
    completed = subprocess.run(
        ["git", "-C", str(source_repo), "show", f"{source_ref}:{source_path}"],
        check=True,
        capture_output=True,
    )
    source_html = completed.stdout
    html_sha256 = hashlib.sha256(source_html).hexdigest()
    if (
        source_ref == SOURCE_COMMIT
        and family_id == "apply-for-school-place"
        and html_sha256 != EXPECTED_SCHOOL_HTML_SHA256
    ):
        raise ValueError(
            "Pinned family HTML identity changed: "
            f"expected {EXPECTED_SCHOOL_HTML_SHA256}, got {html_sha256}"
        )

    source_text = source_html.decode("utf-8")
    matches = re.findall(
        r'<pre id="governed-family-record"><code[^>]*>(.*?)</code></pre>',
        source_text,
        re.DOTALL,
    )
    if len(matches) != 1:
        raise ValueError(
            "Expected exactly one complete governed family record in the source HTML; "
            f"found {len(matches)}"
        )
    record_text = html.unescape(matches[0]).strip()
    record = json.loads(record_text)
    if record.get("schema") != EXPECTED_RECORD_SCHEMA:
        raise ValueError(f"Unexpected record schema: {record.get('schema')!r}")
    actual_projection_sha = record.get("source_projection", {}).get("sha256")
    if actual_projection_sha != EXPECTED_PROJECTION_SHA256:
        raise ValueError(
            "Unexpected source projection SHA-256: "
            f"expected {EXPECTED_PROJECTION_SHA256}, got {actual_projection_sha}"
        )
    actual_family_id = record.get("family", {}).get("id")
    if actual_family_id != family_id:
        raise ValueError(
            f"Family identity mismatch: expected {family_id!r}, got {actual_family_id!r}"
        )
    binding = SourceBinding(
        family_id=family_id,
        source_path=source_path,
        public_url=f"{SOURCE_URL_ROOT}{source_path}",
        source_ref=source_ref,
        html_sha256=html_sha256,
        record_sha256=hashlib.sha256(record_text.encode("utf-8")).hexdigest(),
    )
    return record_text, record, binding


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    colour: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().append(OxmlElement("w:lang"))
    run._element.rPr[-1].set(qn("w:val"), "en-GB")
    if size is not None:
        run.font.size = Pt(size)
    if colour is not None:
        run.font.color.rgb = colour
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name: str, size: float, colour: RGBColor = INK) -> None:
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = colour
    lang = style._element.get_or_add_rPr().find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        style._element.get_or_add_rPr().append(lang)
    lang.set(qn("w:val"), "en-GB")


def configure_document(
    doc: Document,
    *,
    header_label: str = "SharePoint and M365 Copilot experiment",
) -> None:
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, "Calibri", 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, colour, before, after) in heading_specs.items():
        style = doc.styles[name]
        set_style_font(style, "Calibri", size, colour)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Governed JSON" not in doc.styles:
        code_style = doc.styles.add_style("Governed JSON", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = doc.styles["Governed JSON"]
    set_style_font(code_style, "Courier New", 8, INK)
    code_style.paragraph_format.space_before = Pt(0)
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing = 1.0
    code_style.paragraph_format.widow_control = False

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run(f"Explore OKF | {header_label}")
    set_run_font(hr, size=9, colour=MUTED)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fr = fp.add_run("Exploratory governed derivative  |  Page ")
    set_run_font(fr, size=8.5, colour=MUTED)
    add_page_field(fp)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    set_run_font(run, size=8.5, colour=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, text, end):
        run._r.append(node)


def shade_paragraph(paragraph, fill: str, border_colour: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border_colour)
    borders.append(left)
    p_pr.append(borders)


def add_callout(doc: Document, label: str, text: str, *, caution: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.2
    shade_paragraph(p, "FFF1D8" if caution else "E2F4F0", "9C4D00" if caution else "006B62")
    label_run = p.add_run(f"{label} ")
    set_run_font(label_run, bold=True, colour=INK)
    text_run = p.add_run(text)
    set_run_font(text_run, colour=INK)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), "004D47")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([colour, underline])
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def resolve_narrative_link(target: str) -> str | None:
    """Resolve governed narrative links without admitting unsafe schemes."""

    if (
        any(character.isspace() for character in target)
        or any(character in target for character in "\"'<>\\{}|^`")
        or re.search(r"%(?![0-9A-Fa-f]{2})", target)
    ):
        return None
    parsed = urlsplit(target)
    if parsed.scheme in {"http", "https"} and parsed.hostname:
        try:
            port = parsed.port
        except ValueError:
            return None
        if parsed.username is not None or parsed.password is not None:
            return None
        if port is not None and not 1 <= port <= 65_535:
            return None
        return target
    if parsed.scheme or parsed.netloc or target.startswith(("#", "?")):
        return None
    if ".." in Path(target).parts:
        return None
    resolved = urljoin(SOURCE_URL_ROOT, target.lstrip("/"))
    return resolved if resolved.startswith(SOURCE_URL_ROOT) else None


def add_markdown_text(paragraph, text: str) -> None:
    """Render simple governed Markdown links as Word hyperlinks."""

    cleaned = text.replace("**", "")
    cursor = 0
    for match in MARKDOWN_LINK_PATTERN.finditer(cleaned):
        if match.start() > cursor:
            run = paragraph.add_run(cleaned[cursor : match.start()])
            set_run_font(run)
        label, target = match.groups()
        resolved = resolve_narrative_link(target)
        if resolved is None:
            run = paragraph.add_run(label)
            set_run_font(run)
        else:
            add_hyperlink(paragraph, label, resolved)
        cursor = match.end()
    if cursor < len(cleaned):
        run = paragraph.add_run(cleaned[cursor:])
        set_run_font(run)


def add_label_value(
    doc: Document,
    label: str,
    value: Any,
    *,
    code: bool = False,
    url: str | None = None,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, bold=True, colour=INK)
    value_text = str(value)
    if url is not None:
        add_hyperlink(p, value_text, url)
    else:
        value_run = p.add_run(value_text)
        set_run_font(
            value_run,
            name="Courier New" if code else "Calibri",
            size=9.5 if code else 11,
            colour=INK,
        )


def next_abstract_id(numbering) -> int:
    values = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    return max(values, default=-1) + 1


def next_num_id(numbering) -> int:
    values = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    return max(values, default=0) + 1


def create_numbering(doc: Document, *, bullet: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = next_abstract_id(numbering)
    num_id = next_num_id(numbering)

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(level_text)
    alignment = OxmlElement("w:lvlJc")
    alignment.set(qn("w:val"), "left")
    level.append(alignment)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc: Document, text: str, num_id: int, *, code: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend([level, number])
    p_pr.append(num_pr)
    run = p.add_run(text)
    set_run_font(run, name="Courier New" if code else "Calibri", size=9.5 if code else 11)


def add_state_detail(doc: Document, label: str, value: dict[str, Any]) -> None:
    state = value.get("state", "not stated")
    detail = value.get("summary") or value.get("reason") or "No further detail stated."
    add_label_value(doc, label, f"{state} — {detail}")


def add_title_block(doc: Document, record: dict[str, Any]) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(4)
    kr = kicker.add_run("SHAREPOINT AND M365 COPILOT EXPERIMENT")
    set_run_font(kr, size=9.5, colour=TEAL, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    tr = title.add_run(record["family"]["title"])
    set_run_font(tr, size=26, colour=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    sr = subtitle.add_run("Governed family record for SharePoint and AI retrieval testing")
    set_run_font(sr, size=13, colour=MUTED)

    add_callout(
        doc,
        "Exploratory.",
        "This is a deterministic derivative of one governed OKF family record. "
        "It is a discovery aid, not an official service or source of current or personalised advice.",
        caution=True,
    )


def add_gate_block(
    doc: Document,
    record: dict[str, Any],
    binding: SourceBinding,
) -> None:
    doc.add_heading("Retrieval and provenance gate", level=1)
    add_label_value(doc, "Record schema", record["schema"], code=True)
    add_label_value(doc, "Source projection SHA-256", record["source_projection"]["sha256"], code=True)
    add_label_value(doc, "Source projection path", record["source_projection"]["path"], code=True)
    add_label_value(doc, "Source projection bytes", record["source_projection"]["bytes"], code=True)
    add_label_value(doc, "Snapshot", record["snapshot"], code=True)
    add_label_value(doc, "Generated at", record["generated_at"], code=True)
    add_label_value(doc, "Stable family ID", record["family"]["id"], code=True)
    add_label_value(
        doc,
        "Public governed record",
        binding.public_url,
        url=binding.public_url,
    )
    add_label_value(doc, "Source Git commit", binding.source_ref, code=True)
    add_label_value(doc, "Family HTML SHA-256", binding.html_sha256, code=True)
    add_label_value(doc, "Governed record SHA-256", binding.record_sha256, code=True)
    add_callout(
        doc,
        "Gatekeeping rule.",
        "An AI should report the record schema and source projection SHA-256 before using "
        "the family content. If either value is unavailable, it should stop rather than infer details.",
    )


def add_authored_narrative(doc: Document, record: dict[str, Any]) -> None:
    doc.add_heading("Authored narrative", level=1)
    narrative = record["family"]["narrative"]["body"]
    lines = narrative.splitlines()
    paragraph_buffer: list[str] = []
    bullet_id = create_numbering(doc, bullet=True)

    def flush_paragraph() -> None:
        if paragraph_buffer:
            rendered = " ".join(part.strip() for part in paragraph_buffer)
            paragraph = doc.add_paragraph()
            add_markdown_text(paragraph, rendered)
            paragraph_buffer.clear()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            continue
        if stripped.startswith("# "):
            continue
        if stripped.startswith("## "):
            flush_paragraph()
            doc.add_heading(stripped[3:], level=2)
            continue
        if stripped.startswith("- ["):
            flush_paragraph()
            match = re.fullmatch(r"- \[([^]]+)]\(([^)]+)\)", stripped)
            if match:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                p_pr = p._p.get_or_add_pPr()
                num_pr = OxmlElement("w:numPr")
                level = OxmlElement("w:ilvl")
                level.set(qn("w:val"), "0")
                number = OxmlElement("w:numId")
                number.set(qn("w:val"), str(bullet_id))
                num_pr.extend([level, number])
                p_pr.append(num_pr)
                target = resolve_narrative_link(match.group(2))
                if target is None:
                    run = p.add_run(match.group(1))
                    set_run_font(run)
                else:
                    add_hyperlink(p, match.group(1), target)
            else:
                add_list_item(doc, stripped[2:], bullet_id)
            continue
        paragraph_buffer.append(stripped)
    flush_paragraph()


def add_family_identity(doc: Document, family: dict[str, Any]) -> None:
    doc.add_heading("Family identity and coverage", level=1)
    add_label_value(doc, "Title", family["title"])
    add_label_value(doc, "Stable family ID", family["id"], code=True)
    add_label_value(doc, "Family route", family["route"], code=True)
    add_label_value(doc, "Description", family["description"])
    add_label_value(doc, "Domain", f"{family['domain']['title']} ({family['domain']['id']})")
    add_label_value(doc, "Domain route", family["domain"]["route"], code=True)
    add_label_value(doc, "Enclosing process", f"{family['process']['title']} ({family['process']['id']})")
    add_label_value(doc, "Process route", family["process"]["route"], code=True)
    add_label_value(doc, "Population status", family["status"], code=True)
    add_label_value(doc, "Assertion status", family["assertion_status"], code=True)
    add_label_value(doc, "Interaction boundary", family["interaction_boundary"])

    doc.add_heading("Useful terms", level=2)
    bullet_id = create_numbering(doc, bullet=True)
    for alias in family["aliases"]:
        add_list_item(doc, alias, bullet_id)

    doc.add_heading("Example situations", level=2)
    bullet_id = create_numbering(doc, bullet=True)
    for situation in family["situations"]:
        add_list_item(doc, situation, bullet_id)

    doc.add_heading("User needs", level=2)
    bullet_id = create_numbering(doc, bullet=True)
    for need in family["user_needs"]:
        add_list_item(doc, need, bullet_id)


def add_applicability(doc: Document, family: dict[str, Any]) -> None:
    doc.add_heading("Jurisdiction routes", level=1)
    intro = doc.add_paragraph(
        "Each jurisdiction is supported through an explicit route. Similar labels do not mean "
        "that rules, eligibility, evidence, costs, deadlines or remedies are the same."
    )
    intro.paragraph_format.space_after = Pt(10)
    for item in family["applicability"]:
        doc.add_heading(item["jurisdiction"], level=2)
        add_label_value(doc, "Applicability state", item["state"], code=True)
        for variant in item["route_variants"]:
            add_label_value(doc, "Route ID", variant["id"], code=True)
            if "route_kind" in variant:
                add_label_value(doc, "Route kind", variant["route_kind"], code=True)
            add_label_value(doc, "Provider", variant["provider"], code=True)
            add_label_value(doc, "Primary source ID", variant["primary_source"], code=True)


def add_episode(doc: Document, episode: dict[str, Any]) -> None:
    doc.add_heading(f"{episode['title']} ({episode['kind']})", level=2)
    add_label_value(doc, "Episode ID", episode["id"], code=True)
    add_label_value(doc, "Episode order", episode["order"], code=True)
    add_label_value(doc, "Episode route", episode["route"], code=True)
    add_label_value(doc, "Entry state", episode["entry_state"])
    add_label_value(doc, "Episode outcome", episode["outcome"])

    number_id = create_numbering(doc, bullet=False)
    for step in sorted(episode["steps"], key=lambda item: item["order"]):
        add_list_item(doc, step["interaction"], number_id)
        add_label_value(doc, "Step ID", step["id"], code=True)
        add_label_value(doc, "Step order", step["order"], code=True)
        add_label_value(doc, "Step route", step["route"], code=True)
        add_label_value(doc, "Provider", step["provider"], code=True)
        add_label_value(doc, "Assertion status", step["assertion_status"], code=True)
        add_state_detail(doc, "Requirements", step["requirements"])
        add_state_detail(doc, "Evidence", step["evidence"])
        add_state_detail(doc, "Rule", step["rule"])
        add_state_detail(doc, "Channel", step["channel"])
        add_state_detail(doc, "Cost", step["cost"])
        add_state_detail(doc, "Time", step["time"])
        add_state_detail(doc, "Output", step["output"])
        add_state_detail(doc, "Outcome", step["outcome"])
        add_state_detail(doc, "Redress", step["redress"])
        add_label_value(doc, "Source IDs", ", ".join(step["source_ids"]), code=True)


def add_journeys(doc: Document, family: dict[str, Any]) -> None:
    doc.add_heading("Ordered journey", level=1)
    for episode in sorted(family["episodes"], key=lambda item: item["order"]):
        add_episode(doc, episode)


def add_sources(doc: Document, family: dict[str, Any]) -> None:
    doc.add_heading("Official source handoffs", level=1)
    doc.add_paragraph(
        "These links locate current official services. Check the relevant official source before acting."
    )
    for source in sorted(family["sources"], key=lambda item: item["position"]):
        jurisdiction = ", ".join(source["jurisdictions"]) or "Jurisdiction not stated"
        doc.add_heading(f"{jurisdiction}: {source['title']}", level=2)
        add_label_value(doc, "Official URL", source["url"], url=source["url"])
        add_label_value(doc, "Owner", source["owner"])
        add_label_value(doc, "Source ID", source["id"], code=True)
        add_label_value(doc, "Source route", source["route"], code=True)
        add_label_value(doc, "Authority role", source["authority_role"], code=True)
        add_label_value(doc, "Summary", source["summary"])
        add_label_value(doc, "Observed at", source["observed_at"], code=True)
        add_label_value(doc, "Freshness", source["freshness"], code=True)
        add_label_value(doc, "Jurisdiction basis", source["jurisdiction_basis"], code=True)
        add_label_value(doc, "Rights decision", source["rights_decision"], code=True)
        bullet_id = create_numbering(doc, bullet=True)
        for limitation in source["limitations"]:
            add_list_item(doc, limitation, bullet_id)


def add_review_and_related(doc: Document, family: dict[str, Any]) -> None:
    doc.add_heading("Review status and limitations", level=1)
    add_label_value(doc, "Population gate", family["review"]["population_gate"], code=True)
    add_label_value(doc, "Specialist review", family["review"]["specialist_review"], code=True)
    add_callout(
        doc,
        "Important distinction.",
        "Population-complete means the family is present for discovery. It does not mean the "
        "record has passed specialist review or that operational claims are release-grade.",
        caution=True,
    )
    bullet_id = create_numbering(doc, bullet=True)
    for limitation in family["limitations"]:
        add_list_item(doc, limitation, bullet_id)

    doc.add_heading("Related families", level=1)
    for related in family["related_families"]:
        add_label_value(
            doc,
            related["title"],
            f"{related['id']} — {related['relationship']}; sequenced: "
            f"{str(related['sequenced']).lower()}",
            code=True,
        )


def add_technical_appendix(doc: Document, record_text: str) -> None:
    doc.add_page_break()
    doc.add_heading("Complete governed record (technical appendix)", level=1)
    doc.add_paragraph(
        "The following is the complete JSON envelope embedded in the governed family HTML. "
        "It is included so identifiers, assertion links and source bindings remain available "
        "to reviewers and retrieval systems without inference."
    )
    for line in record_text.splitlines():
        p = doc.add_paragraph(style="Governed JSON")
        run = p.add_run(line if line else " ")
        set_run_font(run, name="Courier New", size=8, colour=INK)


def build_document(
    record_text: str,
    record: dict[str, Any],
    binding: SourceBinding,
    output: Path,
) -> None:
    doc = Document()
    configure_document(doc, header_label="Full governed record v1")
    family = record["family"]
    doc.core_properties.title = f"{family['title']} — governed family record"
    doc.core_properties.subject = "SharePoint and Microsoft 365 Copilot retrieval experiment"
    doc.core_properties.author = "OKF Explorer"
    doc.core_properties.last_modified_by = "OKF Explorer"
    doc.core_properties.keywords = (
        "OKF, Open Knowledge Format, SharePoint, Microsoft 365 Copilot, "
        f"{family['id']}, {', '.join(family['aliases'])}"
    )
    doc.core_properties.comments = (
        "Deterministically generated from the pinned OKF UK Living governed family record."
    )
    source_time = datetime.fromisoformat(record["generated_at"].replace("Z", "+00:00"))
    doc.core_properties.created = source_time
    doc.core_properties.modified = source_time
    doc.core_properties.revision = 1

    add_title_block(doc, record)
    add_gate_block(doc, record, binding)
    add_authored_narrative(doc, record)
    add_family_identity(doc, family)
    add_applicability(doc, family)
    add_journeys(doc, family)
    add_sources(doc, family)
    add_review_and_related(doc, family)
    add_technical_appendix(doc, record_text)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    args = parse_args()
    record_text, record, binding = read_governed_record(
        args.source_repo,
        args.source_ref,
        args.family,
    )
    output = args.output
    if output is None:
        output = (
            Path(__file__).resolve().parent
            / args.family
            / f"{args.family}.docx"
        )
    build_document(record_text, record, binding, output)
    print(output)


if __name__ == "__main__":
    main()
